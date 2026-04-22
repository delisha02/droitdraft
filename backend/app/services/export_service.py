import io
import re
from typing import Optional
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from xhtml2pdf import pisa
from bs4 import BeautifulSoup

class ExportService:
    """
    Service to convert HTML content from the editor into professional legal documents.
    """

    def generate_pdf(self, html_content: str, title: str, font: str = "Times New Roman", font_size: str = "12") -> io.BytesIO:
        """
        Generates a PDF from HTML content using xhtml2pdf.
        """
        # Wrap content in a full HTML structure with styling
        styled_html = f"""
        <html>
            <head>
                <style>
                    @page {{
                        size: a4 portrait;
                        margin: 1in;
                    }}
                    body {{
                        font-family: "{font}", serif;
                        font-size: {font_size}pt;
                        line-height: 1.5;
                        color: #333;
                    }}
                    h1 {{ font-size: 18pt; text-align: center; font-weight: bold; margin-bottom: 20px; }}
                    h2 {{ font-size: 14pt; font-weight: bold; margin-top: 15px; margin-bottom: 10px; }}
                    h3 {{ font-size: 12pt; font-weight: bold; margin-top: 10px; }}
                    p {{ margin-bottom: 10px; text-align: justify; }}
                    table {{ width: 100%; border-collapse: collapse; margin-bottom: 20px; }}
                    th, td {{ border: 1px solid #000; padding: 8px; text-align: left; }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
        </html>
        """
        
        pdf_buffer = io.BytesIO()
        pisa_status = pisa.CreatePDF(styled_html, dest=pdf_buffer)
        
        if pisa_status.err:
            raise Exception(f"PDF generation failed: {pisa_status.err}")
            
        pdf_buffer.seek(0)
        return pdf_buffer

    def generate_docx(self, html_content: str, title: str, font: str = "Times New Roman", font_size: str = "12") -> io.BytesIO:
        """
        Generates a DOCX from HTML content by parsing tags and applying docx styles.
        """
        doc = DocxDocument()
        
        # Set default font
        style = doc.styles['Normal']
        style.font.name = font
        style.font.size = Pt(int(font_size))
        
        # Use BeautifulSoup to parse HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        for element in soup.children:
            if element.name == 'h1':
                p = doc.add_heading(element.get_text(), level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif element.name == 'h2':
                doc.add_heading(element.get_text(), level=1)
            elif element.name == 'h3':
                doc.add_heading(element.get_text(), level=2)
            elif element.name == 'p':
                p = doc.add_paragraph()
                self._process_html_inline(element, p)
            elif element.name == 'table':
                self._process_html_table(element, doc)
            elif element.name is None: # Text nodes outside tags
                if element.strip():
                    doc.add_paragraph(element.strip())

        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        return docx_buffer

    def _process_html_inline(self, element, paragraph):
        """Recursively process inline tags like strong, em, br."""
        for child in element.children:
            if child.name == 'strong' or child.name == 'b':
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name == 'em' or child.name == 'i':
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == 'br':
                paragraph.add_run().add_break()
            elif child.name is None: # Plain text
                paragraph.add_run(child)
            else:
                # Handle nested or unknown tags by just taking text
                self._process_html_inline(child, paragraph)

    def _process_html_table(self, table_tag, doc):
        """Process HTML table tags into docx tables."""
        rows = table_tag.find_all('tr')
        if not rows:
            return
            
        # Determine max columns
        max_cols = 0
        for row in rows:
            max_cols = max(max_cols, len(row.find_all(['td', 'th'])))
            
        if max_cols == 0:
            return
            
        docx_table = doc.add_table(rows=len(rows), cols=max_cols)
        docx_table.style = 'Table Grid'
        
        for r_idx, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for c_idx, cell in enumerate(cells):
                if c_idx < max_cols:
                    docx_cell = docx_table.cell(r_idx, c_idx)
                    docx_cell.text = cell.get_text().strip()

export_service = ExportService()
