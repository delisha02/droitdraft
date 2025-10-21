import os
import pytest
import shutil
from docx import Document as DocxDocument
from PIL import Image, ImageDraw, ImageFont
from app.agents.document_processor.docx_processor import extract_text_from_docx
from app.agents.document_processor.pdf_processor import extract_text_from_pdf
from app.agents.document_processor.text_extractor import extract_text
from unittest.mock import MagicMock

@pytest.fixture(scope="module")
def docx_file():
    """
    Creates a dummy .docx file for testing.
    """
    file_path = "test.docx"
    document = DocxDocument()
    document.add_paragraph("This is a test document.")
    document.add_paragraph("This is the second paragraph.")
    document.save(file_path)
    yield file_path
    os.remove(file_path)

@pytest.fixture(scope="module")
def image_file():
    """
    Creates a dummy image file for testing OCR.
    """
    file_path = "test.png"
    img = Image.new('RGB', (400, 100), color = (255, 255, 255))
    d = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 40)
    except IOError:
        font = ImageFont.load_default()
    d.text((10,10), "This is a test", fill=(0,0,0), font=font)
    img.save(file_path)
    yield file_path
    os.remove(file_path)

def is_tesseract_installed():
    """
    Checks if Tesseract is installed.
    """
    return shutil.which("tesseract") is not None

def test_extract_text_from_docx(docx_file):
    """
    Tests the extraction of text from a .docx file.
    """
    text = extract_text_from_docx(docx_file)
    assert "This is a test document." in text
    assert "This is the second paragraph." in text

def test_extract_text_from_pdf(mocker):
    """
    Tests the extraction of text from a text-based PDF file.
    """
    # Mock pdfplumber
    mock_pdf = MagicMock()
    mock_page = MagicMock()
    mock_page.extract_text.return_value = "This is a test PDF document."
    mock_pdf.pages = [mock_page]
    
    # The __enter__ and __exit__ methods need to be mocked for the 'with' statement
    mock_open = mocker.patch("pdfplumber.open")
    mock_open.return_value.__enter__.return_value = mock_pdf

    text = extract_text_from_pdf("dummy.pdf")
    assert "This is a test PDF document." in text

@pytest.mark.skipif(not is_tesseract_installed(), reason="Tesseract is not installed")
def test_extract_text_from_image(image_file):
    """
    Tests the extraction of text from an image file using OCR.
    """
    from app.agents.document_processor.ocr_processor import extract_text_from_image
    text = extract_text_from_image(image_file)
    assert "This is a test" in text

def test_extract_text_unified(docx_file, mocker, image_file):
    """
    Tests the unified extract_text function.
    """
    text_docx = extract_text(docx_file)
    assert "This is a test document." in text_docx

    # Mock pdfplumber for the PDF test
    mock_pdf = MagicMock()
    mock_page = MagicMock()
    mock_page.extract_text.return_value = "This is a test PDF document."
    mock_pdf.pages = [mock_page]
    
    mock_open = mocker.patch("pdfplumber.open")
    mock_open.return_value.__enter__.return_value = mock_pdf

    text_pdf = extract_text("dummy.pdf")
    assert "This is a test PDF document." in text_pdf

    if is_tesseract_installed():
        text_image = extract_text(image_file)
        assert "This is a test" in text_image